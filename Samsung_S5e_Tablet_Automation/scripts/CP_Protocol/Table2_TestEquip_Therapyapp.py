import os
import time
import subprocess
from appium import webdriver
from appium.options.android import UiAutomator2Options
from appium.webdriver.common.appiumby import AppiumBy as By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.common.exceptions import TimeoutException, NoSuchElementException
from selenium.webdriver.common.actions.pointer_input import PointerInput
from selenium.webdriver.common.actions.interaction import Interaction
from selenium.webdriver.common.action_chains import ActionBuilder
from selenium.webdriver import ActionChains
from selenium.webdriver.common.actions import interaction
from PIL import Image
from docx import Document
from docx.shared import Inches
import sys

sys.stdout.reconfigure(encoding="utf-8")

# add project root
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))

# --------------------------------------------------
# CORE IMPORTS
# --------------------------------------------------
from core.device_registry import is_device_allowed, get_device_name
from core.adb_utils import (
    get_connected_device_serial,
    is_device_locked,
    unlock_device_with_password
)
from core.config_loader import load_device_password
from datetime import datetime

def log(section, message):
    ts = datetime.now().strftime("%H:%M:%S")
    print(f"[{ts}] [{section:<14}] {message}")

# --------------------------------------------------
# DEVICE RESOLUTION & UNLOCK
# --------------------------------------------------
device_serial = os.environ.get("DEVICE_SERIAL") or get_connected_device_serial()

if not is_device_allowed(device_serial):
    raise RuntimeError(f"Unauthorized S5E device connected: {device_serial}")

device_name = get_device_name(device_serial)
# print(f"Authorized device detected: {device_name} ({device_serial})")
log("DEVICE", f"Authorized device detected: {device_name} ({device_serial})")


if is_device_locked():
    unlock_device_with_password(load_device_password())
else:
    # print("Device already unlocked.")
    log("DEVICE", "Device already unlocked")
    
    
# --------------------------------------------------
# START APPIUM
# --------------------------------------------------
from appiumManager.appium_manager import AppiumManager
AppiumManager.start()


# ====== HubCode1 Functions ======
def tap_by_bounds(driver, bounds_str):
    bounds = bounds_str.replace("][", ",").replace("[", "").replace("]", "").split(",")
    x1, y1, x2, y2 = map(int, bounds)
    center_x = (x1 + x2) // 2
    center_y = (y1 + y2) // 2

    finger = PointerInput(Interaction.TOUCH, "finger")
    action = ActionBuilder(driver, mouse=finger)
    action.pointer_action.move_to_location(center_x, center_y)
    action.pointer_action.pointer_down()
    action.pointer_action.pause(0.1)
    action.pointer_action.pointer_up()
    action.perform()

def force_stop_and_launch_hub(serial):
    # print(f"Launching Hub on {serial}")
    log("Hub", f"Launching Hub on {serial}")

    subprocess.run(
        f"adb -s {serial} shell am force-stop com.airwatch.androidagent",
        shell=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL
    )

    time.sleep(2)

    subprocess.run(
        f"adb -s {serial} shell monkey -p com.airwatch.androidagent "
        f"-c android.intent.category.LAUNCHER 1",
        shell=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL
    )

    time.sleep(5)

def force_stop_and_launch_settings(serial):
    # print(f"Launching Settings on {serial}")
    log("SYSTEM UPDATE", f"Launching Settings on {serial}")

    subprocess.run(
        f"adb -s {serial} shell am force-stop com.android.settings",
        shell=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL
    )

    time.sleep(2)

    subprocess.run(
        f"adb -s {serial} shell monkey -p com.android.settings "
        f"-c android.intent.category.LAUNCHER 1",
        shell=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL
    )

    time.sleep(5)

def stop_appium_server():
    log("APPIUM", "Stopping Appium server")
    subprocess.run(
        "taskkill /F /IM node.exe /T",
        shell=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL
    )
    time.sleep(3)


def start_appium_server():
    log("APPIUM", "Starting Appium server")
    subprocess.Popen(
        "appium -p 4723",
        shell=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL
    )
    time.sleep(10)

def get_driver_for_hub(serial):
    options = UiAutomator2Options()
    options.platform_name = "Android"
    options.device_name = serial
    options.udid = serial
    options.app_package = "com.airwatch.androidagent"
    options.app_activity = "com.airwatch.androidagent.AgentHomeActivity"
    options.no_reset = True
    options.new_command_timeout = 300

    return webdriver.Remote("http://127.0.0.1:4723", options=options)

def wait_for_package(driver, package_name="com.airwatch.androidagent", timeout=20):
    print(f"Waiting for package '{package_name}' to be active...")
    for _ in range(timeout):
        if driver.current_package == package_name:
            print(f"Package detected: {package_name}")
            return True
        time.sleep(1)
    print(f"Timeout waiting for package {package_name}")
    return False

def wait_for_hub_screen(driver, timeout=20):
    # print("Waiting for Hub UI (EditText)...")
    log("HUB", "Waiting for Hub UI")
    try:
        WebDriverWait(driver, timeout).until(
            EC.presence_of_element_located((By.CLASS_NAME, "android.widget.EditText"))
        )
        print("Hub UI detected.")
    except TimeoutException:
        print("Timeout waiting for Hub UI.")
        log("HUB", "Timeout waiting for Hub UI.")

# ====== HubCode2 Functions ======

def swipe_up_left_side(driver):
    size = driver.get_window_size()
    start_x = size['width'] * 0.2
    start_y = size['height'] * 0.8
    end_y = size['height'] * 0.2
    actions = ActionChains(driver)
    pointer = PointerInput(interaction.POINTER_TOUCH, "finger")
    actions.w3c_actions = ActionBuilder(driver, mouse=pointer)
    actions.w3c_actions.pointer_action.move_to_location(start_x, start_y)
    actions.w3c_actions.pointer_action.pointer_down()
    actions.w3c_actions.pointer_action.move_to_location(start_x, end_y)
    actions.w3c_actions.pointer_action.release()
    actions.perform()
    time.sleep(1)

def swipe_up_right_side(driver):
    size = driver.get_window_size()
    start_x = size['width'] * 0.8
    start_y = size['height'] * 0.85
    end_y = size['height'] * 0.15

    actions = ActionChains(driver)
    pointer = PointerInput(interaction.POINTER_TOUCH, "finger")
    actions.w3c_actions = ActionBuilder(driver, mouse=pointer)

    actions.w3c_actions.pointer_action.move_to_location(start_x, start_y)
    actions.w3c_actions.pointer_action.pointer_down()
    actions.w3c_actions.pointer_action.pause(0.4)
    actions.w3c_actions.pointer_action.move_to_location(start_x, end_y)
    actions.w3c_actions.pointer_action.release()

    actions.perform()
    time.sleep(1)

def scroll_to_top_left(driver, max_swipes=3):
    for _ in range(max_swipes):
        swipe_up_left_side(driver)
        time.sleep(1)

def find_and_click_about_tablet(driver):
    try:
        element = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//android.widget.TextView[@text="About tablet"]'))
        )
        element.click()
        log("SYSTEM UPDATE", "Clicking on About tablet")

    except:
        scroll_to_top_left(driver)
        find_and_click_about_tablet(driver)

def find_and_click_software_info(driver):
    try:
        element = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//android.widget.TextView[@resource-id="android:id/title" and @text="Software information"]'))
        )
        element.click()
        # print("Clicked on 'Software information'")
        log("SYSTEM UPDATE", "Opening Software information")
    except:
        print("Retrying to find 'Software information'...")
        swipe_up_left_side(driver)
        find_and_click_software_info(driver)

def find_and_click_google_play_system_update(driver):
    try:
        element = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((
                By.XPATH,
                '//androidx.recyclerview.widget.RecyclerView[@resource-id="com.android.settings:id/recycler_view"]'
                '/android.widget.LinearLayout[3]/android.widget.RelativeLayout'
            ))
        )
        element.click()
        log("SYSTEM UPDATE", "Opening Google Play system update")
        return True
    except:
        log("SYSTEM UPDATE", "Google Play system update not found")
        return False

def get_google_play_patch_from_software_info(driver):
    try:
        elements = driver.find_elements(By.CLASS_NAME, "android.widget.TextView")

        for i, el in enumerate(elements):
            if el.text.strip() == "Google Play system update":
                if i + 1 < len(elements):
                    return elements[i + 1].text.strip()

        return "N/A"
    except Exception as e:
        log("SETTINGS", f"Failed to read Google Play update: {e}")
        return "N/A"


def handle_google_play_update(driver):
    try:
        max_cycles = 40
        cycles = 0

        while cycles < max_cycles:
            cycles += 1
            time.sleep(2)

            elements = driver.find_elements(By.CLASS_NAME, "android.widget.TextView")
            texts = [el.text.strip() for el in elements if el.text.strip()]
            print("Visible Texts:", texts)

            # CASE 0 — RESTART DETECTED (EXIT IMMEDIATELY)
            for text in texts:
                if "restart" in text.lower():
                    log("SETTINGS", "Restart required — exiting update handler")
                    return "RESTART_REQUIRED"

            # CASE 1 — Extract version if visible
            for i, text in enumerate(texts):
                if "Google Play system update:" in text:
                    patch = text.split("Google Play system update:", 1)[1].strip()
                    log("SETTINGS", f"Google Play update found: {patch}")
                    return patch

                if text == "Google Play system update" and i + 1 < len(texts):
                    patch = texts[i + 1].strip()
                    log("SETTINGS", f"Google Play update found: {patch}")
                    return patch

            # CASE 2 — Try click update button
            try:
                btn = WebDriverWait(driver, 5).until(
                    EC.presence_of_element_located((
                        By.XPATH,
                        '//android.widget.Button[@resource-id="com.android.vending:id/0_resource_name_obfuscated"]'
                    ))
                )

                btn_text = btn.text.strip()
                log("SETTINGS", f"Found update button: {btn_text if btn_text else 'Unknown'}")
                btn.click()

                # IF RESTART IS TRIGGERED → EXIT, DO NOT WAIT
                if "restart" in btn_text.lower():
                    log("SETTINGS", "Restart initiated — exiting update handler")
                    return "RESTART_INITIATED"

                # Download only — safe to wait
                if "download" in btn_text.lower():
                    log("SETTINGS", "Download started — waiting 5 minutes")
                    time.sleep(300)

                time.sleep(5)
                continue

            except Exception as e:
                log("SETTINGS", f"Update button not detected yet: {e}")

            # CASE 3 — REMOVE FORCE TAP (DANGEROUS DURING SYSTEM UI)
            # Force-tapping during system update screens can crash UiAutomator2
            # → intentionally removed

            log("SETTINGS", "Waiting for update UI...")
            time.sleep(3)

        log("SETTINGS", "Timeout waiting for Google Play update")
        return "NO_UPDATE"

    except Exception as e:
        log("ERROR", f"Google Play update handler failed: {e}")
        return "FAILED"

def is_on_google_play_update_screen(driver):
    texts = [
        el.text.strip().lower()
        for el in driver.find_elements(By.CLASS_NAME, "android.widget.TextView")
        if el.text
    ]
    return any("google play system update" in t for t in texts)

def create_base_driver(serial):
    options = UiAutomator2Options()
    options.platform_name = "Android"
    options.device_name = serial
    options.udid = serial
    options.automation_name = "UiAutomator2"

    # BIND TO SETTINGS
    options.app_package = "com.android.settings"
    options.app_activity = "com.android.settings.Settings"

    options.no_reset = True

    # Samsung stability
    options.allow_running_instrumentation = True
    options.ignore_hidden_api_policy_error = True
    
    options.set_capability("skipServerInstallation", True) 
    options.set_capability("skipDeviceInitialization", True)

    options.uiautomator2_server_launch_timeout = 180000
    options.uiautomator2_server_install_timeout = 180000
    options.adb_exec_timeout = 180000
    options.android_install_timeout = 180000

    return webdriver.Remote("http://127.0.0.1:4723", options=options)

def wait_for_settings_ready(driver, timeout=30):

    log("SYSTEM UPDATE", "Waiting for Settings UI to be ready")

    try:
        WebDriverWait(driver, timeout).until(
            EC.presence_of_element_located(
                (By.XPATH, "//android.widget.TextView")
            )
        )

        time.sleep(2)  # allow Samsung UI animation

        log("SYSTEM UPDATE", "Settings UI ready")

    except TimeoutException:
        log("SYSTEM UPDATE", "Settings UI did not fully load — continuing anyway")

def run_Google_playstore_update(serial):

    force_stop_and_launch_settings(serial)
    time.sleep(5)

    driver = create_base_driver(serial)

    wait_for_settings_ready(driver)

    try:
        scroll_to_top_left(driver)
        find_and_click_about_tablet(driver)
        find_and_click_software_info(driver)
        time.sleep(2)

        google_play_patch = get_google_play_patch_from_software_info(driver)
        log("SETTINGS", f"Google Play system update: {google_play_patch}")

        # Screenshot folder
        screenshots_dir = os.path.join(os.path.expanduser("~"), "Desktop", "Appium_Screenshots")
        os.makedirs(screenshots_dir, exist_ok=True)

        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")

        top_path = os.path.join(screenshots_dir, f"Software_Info_TOP_{timestamp}.png")
        bottom_path = os.path.join(screenshots_dir, f"Software_Info_BOTTOM_{timestamp}.png")

        # Screenshot TOP
        time.sleep(2)
        driver.get_screenshot_as_file(top_path)

        # Scroll and screenshot BOTTOM
        swipe_up_right_side(driver)
        time.sleep(2)
        driver.get_screenshot_as_file(bottom_path)

        return google_play_patch, top_path, bottom_path

    finally:
        driver.quit()


def get_enrollment_value(driver, label_text):
    """
    Reads Enrollment value by label text from Hub Enrollment screen
    """
    try:
        value = driver.find_element(
            By.XPATH,
            f"//android.widget.TextView[@text='{label_text}']"
            "/following-sibling::android.widget.TextView"
        )
        return value.text.strip()
    except Exception:
        return "N/A"

def extract_uat_and_group_id(serial):
    # print("Launching Hub app and extracting info...")
    log("Hub", "Launching Hub app and extracting info...")

    force_stop_and_launch_hub(serial)
    time.sleep(8)

    options = UiAutomator2Options()
    options.platform_name = "Android"
    options.device_name = serial
    options.automation_name = "UiAutomator2"
    # options.app_package = "com.airwatch.androidagent"
    # options.app_activity = "com.airwatch.agent.Hub.hostactivity.HostActivity"
    options.no_reset = True

    # CRITICAL FIX
    options.dont_stop_app_on_reset = True
    #options.skip_server_installation = True
    options.set_capability("skipServerInstallation", True)
    options.set_capability("skipDeviceInitialization", True)

    # REQUIRED for Samsung Hub after Settings
    options.uiautomator2_server_launch_timeout = 180000
    options.uiautomator2_server_install_timeout = 180000
    options.adb_exec_timeout = 180000
    options.android_install_timeout = 180000

    options.allow_running_instrumentation = True
    options.ignore_hidden_api_policy_error = True

    driver = webdriver.Remote("http://127.0.0.1:4723", options=options)
    wait = WebDriverWait(driver, 30)

    try:
        profile_icon = wait.until(EC.element_to_be_clickable((By.ID, "com.airwatch.androidagent:id/user_initials_tv")))
        profile_icon.click()
        log("Hub", "Clicking on Hub Profile")
        time.sleep(2)

        device_entry = wait.until(EC.element_to_be_clickable((By.ID, "com.airwatch.androidagent:id/device_tv")))
        device_entry.click()
        log("Hub", "Clicking on This Device")
        time.sleep(2)

        enrollment = wait.until(EC.element_to_be_clickable((By.XPATH, "//android.widget.TextView[@text='Enrollment']")))
        enrollment.click()
        log("Hub", "Clicking on Enrollment")
        time.sleep(2)

        wait.until(EC.presence_of_all_elements_located((By.ID, "com.airwatch.androidagent:id/listview_value")))

        enrolled_server = get_enrollment_value(driver, "Enrolled Server")
        enrolled_group = get_enrollment_value(driver, "Enrolled Group ID")
        username = get_enrollment_value(driver, "Username")

        # LOG EVERYTHING
        log("HUB", f"Enrolled Server : {enrolled_server}")
        log("HUB", f"Enrolled Group  : {enrolled_group}")
        log("HUB", f"Username        : {username}")

        return enrolled_server, enrolled_group, username

    except Exception as e:
        # print(f"Error during UAT extraction: {e}")
        log("HUB", f"Error during enrollment extraction: {e}")
        return "N/A", "N/A", "N/A"
    finally:
        clear_recent_apps_ui(driver)
        driver.quit()
        log("HUB", "Hub session closed")

def get_device_info(serial):
    def get_value(command):
        try:
            return subprocess.check_output(command, shell=True).decode().strip()
        except Exception as e:
            return f"Error: {e}"

    model = get_value(f"adb -s {serial} shell getprop ro.product.model")
    android_version = get_value(f"adb -s {serial} shell getprop ro.build.version.release")
    patch_level = get_value(f"adb -s {serial} shell getprop ro.build.version.security_patch")
    return {
        "serial": serial,
        "model": model,
        "android_version": android_version,
        "patch_level": patch_level
    }

def wait_for_device_unlock_and_restart_appium(serial):
    log("DEVICE", "Waiting for device to reboot")

    subprocess.run(
        f"adb -s {serial} wait-for-device",
        shell=True
    )

    # Samsung boot stabilization
    time.sleep(60)

    log("DEVICE", "Checking lock state after reboot")

    for _ in range(3):
        if is_device_locked():
            log("DEVICE", "Device locked — unlocking")
            unlock_device_with_password(load_device_password())
            time.sleep(5)
        else:
            break

    stop_appium_server()
    start_appium_server()

    log("DEVICE", "Device + Appium ready")

def fill_table_for_s5e(doc_path, info, top_screenshot=None, bottom_screenshot=None):
    doc = Document(doc_path)

    for table in doc.tables:
        for row in table.rows:
            left_cell = row.cells[0].text.strip()
            right_cell = row.cells[1]

            if "CT900E" in left_cell:
                # Clear cell
                right_cell.text = ""

                # Add text
                p = right_cell.paragraphs[0]
                p.add_run(
                    f"Model: {info['model']}\n"
                    f"SN Number: {info['serial']}\n"
                    f"Android OS Version: {info['android_version']}\n"
                    f"UAT Server: {info['uat']}\n"
                    f"Group ID: {info['group_id']}\n"
                    f"Android Security Patch Level: {info['patch_level']}\n"
                    f"Google play system update: {info['google_play_patch']}\n"
                )

                # TOP screenshot
                if top_screenshot and os.path.exists(top_screenshot):
                    p1 = right_cell.add_paragraph()
                    p1.add_run().add_picture(top_screenshot, width=Inches(4.5))

                # BOTTOM screenshot
                if bottom_screenshot and os.path.exists(bottom_screenshot):
                    p2 = right_cell.add_paragraph()
                    p2.add_run().add_picture(bottom_screenshot, width=Inches(4.5))

                doc.save(doc_path)
                log("DOCUMENT", "Updated document with TOP & BOTTOM screenshots")
                log("DOCUMENT", f"Protocol file updated at: {doc_path}")
                return

# --------------------------------------------------
# CLEANUP
# --------------------------------------------------
# def clear_recent_apps_ui(serial):
    # options = UiAutomator2Options()
    # options.platform_name = "Android"
    # options.device_name = serial
    # options.no_reset = True   # IMPORTANT
    # # DO NOT set app_package / app_activity

    # driver = webdriver.Remote("http://127.0.0.1:4723", options=options)

    # try:
        # # print("Opening Recent Apps…")
        # log("CLEANUP", "Opening Recent Apps")
        # driver.press_keycode(187)
        # time.sleep(3)

        # clear_all = WebDriverWait(driver, 10).until(
            # EC.element_to_be_clickable((By.XPATH, "//*[@text='Close all' or @text='Clear all']"))
        # )
        # clear_all.click()
        # # print("Recent apps cleared via UI.")
        # log("CLEANUP", "Recent apps cleared")

    # except Exception as e:
        # print(f"No recent apps to clear or UI not found: {e}")

    # finally:
        # driver.quit()
        

def clear_recent_apps_ui(driver):

    try:
        log("CLEANUP", "Opening Recent Apps")
        driver.press_keycode(187)
        time.sleep(3)
        clear_all = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable(
                (By.XPATH, "//*[@text='Close all' or @text='Clear all']")
            )
        )

        clear_all.click()
        log("CLEANUP", "Recent apps cleared")

    except Exception as e:
        print(f"No recent apps to clear or UI not found: {e}")

# ====== Main Integration ======

def main():
    # print("Starting Full Integration Script")
    log("INIT", "Starting full integration script")

    serial = get_connected_device_serial()
    if not serial:
        print("No device connected.")
        return

    # print(f"Device Detected: {serial}")
    log("DEVICE", f"Device detected: {serial}")

    # STEP 1: FORCE-LAUNCH SETTINGS & GOOGLE PLAY UPDATE
    # google_play_patch = run_Google_playstore_update(serial)
    # google_play_patch, software_info_screenshot = run_Google_playstore_update(serial)
    google_play_patch, software_top, software_bottom = run_Google_playstore_update(serial)

    # Step 2: Extract UAT and Group ID (HubCode2)
    uat, group_id, username = extract_uat_and_group_id(serial)

    # Step 3: Get device info and update dictionary
    info = get_device_info(serial)
    info['google_play_patch'] = google_play_patch
    info['uat'] = uat
    info['group_id'] = group_id

    # Step 4: Update Word document
    doc_path = r"C:\Users\SVC-Systems-TestPC\OneDrive - Medtronic PLC\Protocol\NDHF1500-221664_CP_Latest_Protocol.docx"
    # doc_path = r"C:\Users\mandat3\OneDrive - Medtronic PLC\Desktop\Appium\App Ver Auto Fill\NDHF1500-221664_Jan_2026_Protocol.docx"
    fill_table_for_s5e(doc_path, info)
    # fill_table_for_s5e(doc_path, info, software_info_screenshot)
    fill_table_for_s5e(
        doc_path,
        info,
        software_top,
        software_bottom
    )

    #clear_recent_apps_ui(device_serial)

    # print("All steps complete.")
    log("RESULT", "Automation completed successfully")


if __name__ == "__main__":
    main()
