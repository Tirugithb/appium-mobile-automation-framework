from appium import webdriver
from appium.options.android import UiAutomator2Options
# from selenium.webdriver.common.by import By
from appium.webdriver.common.appiumby import AppiumBy as By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver import ActionChains
from selenium.webdriver.common.action_chains import ActionBuilder
from selenium.webdriver.common.actions import interaction
from selenium.webdriver.common.actions.pointer_input import PointerInput
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from datetime import datetime
import time
import pandas as pd
import os
import subprocess
import sys
import re

sys.stdout.reconfigure(encoding="utf-8")

# add project root
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))

# --------------------------------------------------
# CORE IMPORTS
# --------------------------------------------------
from core.device_registry import is_device_allowed, get_device_name
from core.adb_utils import (
    get_connected_device_serial,
    is_device_locked,
    unlock_device_with_password
)
from core.config_loader import load_device_password

# --------------------------------------------------
# DISPLAY SERIAL MAPPING
# --------------------------------------------------

DISPLAY_SERIAL_MAP = {
    "22a9e8a977239f73": "R52M10M43NV"
}

# --------------------------------------------------
# LOGGING
# --------------------------------------------------
def log(section, message):
    ts = datetime.now().strftime("%H:%M:%S")
    print(f"[{ts}] [{section:<14}] {message}")

def get_actual_device_serial(adb_id):

    try:

        result = subprocess.check_output(
            f'adb -s {adb_id} shell getprop ro.serialno',
            shell=True
        ).decode().strip()

        return result if result else adb_id

    except:

        return adb_id

# --------------------------------------------------
# DEVICE RESOLUTION & UNLOCK
# --------------------------------------------------
device_serial = os.environ.get("DEVICE_SERIAL") or get_connected_device_serial()

display_serial = DISPLAY_SERIAL_MAP.get(
    device_serial,
    device_serial
)

if not is_device_allowed(device_serial):
    raise RuntimeError(f"Unauthorized S2 device connected: {device_serial}")

device_name = get_device_name(device_serial)
# log("DEVICE", f"Authorized device detected: {device_name} ({device_serial})")
log(
    "DEVICE",
    f"Authorized device detected: {device_name} ({display_serial})"
)
if is_device_locked():
    unlock_device_with_password(load_device_password())
else:
    log("DEVICE", "Device already unlocked")

# --------------------------------------------------
# START APPIUM
# --------------------------------------------------
from appiumManager.appium_manager import AppiumManager

AppiumManager.start()


def force_stop_and_launch_settings(serial):
    # print(f"Launching Settings on {serial}")
    log("SYSTEM UPDATE", f"Launching Settings on {serial}")

    subprocess.run(
        f"adb -s {serial} shell am force-stop com.android.settings",
        shell=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL
    )

    time.sleep(2)

    subprocess.run(
        f"adb -s {serial} shell monkey -p com.android.settings "
        f"-c android.intent.category.LAUNCHER 1",
        shell=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL
    )

    time.sleep(5)


# ========== Appium Setup ==========
def get_settings_driver(serial):
    options = UiAutomator2Options()
    options.platform_name = "Android"
    options.device_name = serial
    options.udid = serial
    options.automation_name = "UiAutomator2"
    options.app_package = "com.android.settings"
    options.app_activity = "com.android.settings.Settings"
    options.no_reset = True
    # ADD – required for Samsung / slow instrumentation
    options.uiautomator2_server_launch_timeout = 180000
    options.uiautomator2_server_install_timeout = 180000
    options.adb_exec_timeout = 180000
    options.android_install_timeout = 180000
    # ADD – because Settings is launched via adb monkey
    options.allow_running_instrumentation = True
    options.ignore_hidden_api_policy_error = True

    # return webdriver.Remote("http://localhost:4723", options=options)
    return webdriver.Remote("http://127.0.0.1:4723", options=options)


# ========== Helper Functions ==========

def swipe_up_left_side(driver):
    size = driver.get_window_size()
    start_x = size['width'] * 0.2
    start_y = size['height'] * 0.8
    end_y = size['height'] * 0.2
    actions = ActionChains(driver)
    pointer = PointerInput(interaction.POINTER_TOUCH, "finger")
    actions.w3c_actions = ActionBuilder(driver, mouse=pointer)
    actions.w3c_actions.pointer_action.move_to_location(start_x, start_y)
    actions.w3c_actions.pointer_action.pointer_down()
    actions.w3c_actions.pointer_action.move_to_location(start_x, end_y)
    actions.w3c_actions.pointer_action.release()
    actions.perform()
    time.sleep(1)


def swipe_up_in_apps_list(driver):
    size = driver.get_window_size()
    start_x = size['width'] * 0.5
    start_y = size['height'] * 0.7  # Start slightly higher
    end_y = size['height'] * 0.5  # Shorter swipe to avoid skipping
    actions = ActionChains(driver)
    pointer = PointerInput(interaction.POINTER_TOUCH, "finger")
    actions.w3c_actions = ActionBuilder(driver, mouse=pointer)
    actions.w3c_actions.pointer_action.move_to_location(start_x, start_y)
    actions.w3c_actions.pointer_action.pointer_down()
    actions.w3c_actions.pointer_action.pause(0.5)
    actions.w3c_actions.pointer_action.move_to_location(start_x, end_y)
    actions.w3c_actions.pointer_action.pause(0.5)
    actions.w3c_actions.pointer_action.release()
    actions.perform()
    time.sleep(3)


def wait_for_apps_list_loaded(driver, timeout=10):
    try:
        WebDriverWait(driver, timeout).until(
            EC.presence_of_element_located((By.XPATH, '//android.widget.TextView[contains(@text, "Apps")]'))
        )
    except:
        print("Apps list did not load.")


def find_and_click_apps(driver):
    try:
        apps_option = WebDriverWait(driver, 20).until(
            EC.presence_of_element_located((By.XPATH, '//android.widget.TextView[@text="Apps"]'))
        )
        apps_option.click()
    except:
        # scroll_to_top_left(driver)
        find_and_click_apps(driver)


def scroll_until_app_found(driver, app_name, max_swipes=6):
    swipes = 0
    previous_page_source = ""
    while swipes < max_swipes:
        try:
            wait_for_apps_list_loaded(driver)
            app_option = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.XPATH, f'//android.widget.TextView[contains(@text, "{app_name}")]'))
            )
            app_option.click()

            # Wait for App Info screen to load
            WebDriverWait(driver, 10).until(
                EC.presence_of_element_located((
                    By.XPATH,
                    '//android.widget.TextView'
                ))
            )

            time.sleep(2)

            return True

        except:
            swipe_up_in_apps_list(driver)
            time.sleep(2)
            current_page_source = driver.page_source
            if current_page_source == previous_page_source:
                break
            previous_page_source = current_page_source
            swipes += 1
    return False


def scroll_up_to_version_info(driver, max_swipes=4):

    version_pattern = re.compile(r'\b\d+(?:\.\d+)+(?:[-\w\[\]\s\.]*)')

    try:

        elements = driver.find_elements(
            By.XPATH,
            '//android.widget.TextView'
        )

        for el in elements:

            text = el.text.strip()

            # Skip unwanted texts
            if any(skip in text.lower() for skip in [
                "mb", "gb", "kb", "used", "storage", "since"
            ]):
                continue

            # Case 1: Explicit "Version"
            if "version" in text.lower():

                clean_text = text.lower().replace("version", "").strip()

                match = version_pattern.search(clean_text)

                if match:
                    return match.group()

            # Case 2: Pure version pattern
            match = version_pattern.search(text)

            if match:
                return match.group()

    except:
        pass

    return "N/A"


def navigate_back(driver):
    try:
        back_button = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//android.widget.ImageButton[@content-desc="Navigate up"]'))
        )
        back_button.click()
        time.sleep(2)
    except:
        print("Back button not found.")


def scroll_to_top_left(driver, max_swipes=3):
    for _ in range(max_swipes):
        swipe_up_left_side(driver)
        time.sleep(2)

# ==================================================
# SCROLL SETTINGS
# ==================================================

def scroll_settings(settings_driver):

    try:

        settings_driver.find_element(
            By.ANDROID_UIAUTOMATOR,
            'new UiScrollable(new UiSelector().scrollable(true)).scrollForward()'
        )

        log("SOFTWARE", "Settings screen scrolled")

        time.sleep(2)

    except Exception:

        log("SOFTWARE", "Settings scroll skipped")


# ==================================================
# HANDLE SCHEDULED SOFTWARE UPDATES
# ==================================================

def handle_schedule_software_updates(settings_driver):

    try:

        log("SOFTWARE", "Checking Scheduled software updates toggle")

        WebDriverWait(settings_driver, 10).until(
            EC.presence_of_element_located((
                By.XPATH,
                "//*[contains(@text,'Scheduled software updates')]"
            ))
        )

        switches = settings_driver.find_elements(
            By.CLASS_NAME,
            "android.widget.Switch"
        )

        if not switches:

            log("SOFTWARE", "Scheduled software updates switch not found")
            return

        switch = switches[0]

        checked = (
            switch.get_attribute("checked") or ""
        ).lower()

        log("SOFTWARE", f"Switch checked state = {checked}")

        # ENABLED
        if checked == "true":

            log("SOFTWARE", "Scheduled software updates is ENABLED")

            switch.click()

            time.sleep(2)

            log("SOFTWARE", "Scheduled software updates disabled")

        # DISABLED
        else:

            log("SOFTWARE", "Scheduled software updates already disabled")

    except Exception as e:

        log("SOFTWARE", f"Scheduled software updates check failed: {e}")


# ==================================================
# SOFTWARE UPDATE CHECK
# ==================================================

def check_software_update():

    try:

        settings_driver = get_settings_driver(device_serial)

        log("SOFTWARE", "Opening Settings")

        # --------------------------------------------------
        # SCROLL SETTINGS
        # --------------------------------------------------

        scroll_settings(settings_driver)

        # --------------------------------------------------
        # OPEN SOFTWARE UPDATE
        # --------------------------------------------------

        log("SOFTWARE", "Opening Software Update")

        software_update = WebDriverWait(settings_driver, 15).until(
            EC.element_to_be_clickable((
                By.XPATH,
                "//*[contains(@text,'Software update')]"
            ))
        )

        software_update.click()

        time.sleep(5)

        # --------------------------------------------------
        # HANDLE TOGGLE
        # --------------------------------------------------

        handle_schedule_software_updates(settings_driver)


        settings_driver.quit()

    except Exception as e:

        log("SOFTWARE", f"Software update check failed: {e}")


# ========== Start Process ==========
force_stop_and_launch_settings(device_serial)
driver = get_settings_driver(device_serial)

scroll_to_top_left(driver)
find_and_click_apps(driver)

apps_to_find = ["Adobe Acrobat", "Brother Print Service Plugin", "Calculator", "Camera", "Canon Print Service", "Chrome",
                "Clock", "Content", "Epson iPrint", "Gallery", "Google Play Store", "HP Print Service Plugin", "Hub",
                "Mopria Print Service", "My Files", "Samsung Print Service Plugin", "Web", "Xerox Print Service", "Zoom"]

# apps_to_find = ["Adobe Acrobat", "Brother Print Service Plugin", "Calculator"]

app_data = []

for app_name in apps_to_find:
    if scroll_until_app_found(driver, app_name):
        version_info = scroll_up_to_version_info(driver)
        version_number = version_info.replace("Version ", "") if version_info else "N/A"
        # print(f"{app_name} : {version_number}")
        log("SETTINGS APPS", f"{app_name:<30} : {version_number}")
        app_data.append({"App": app_name, "Version": version_number})
        navigate_back(driver)
    else:
        # print(f"{app_name} not found.")
        log("SETTINGS APPS", f"{app_name:<30} : N/A")
        app_data.append({"App": app_name, "Version": "N/A"})

driver.quit()


# --------------------------------------------------
# CLEANUP
# --------------------------------------------------
def clear_recent_apps_ui(serial):
    options = UiAutomator2Options()
    options.platform_name = "Android"
    options.device_name = serial
    options.no_reset = True  # IMPORTANT
    # DO NOT set app_package / app_activity

    # driver = webdriver.Remote("http://localhost:4723", options=options)
    driver = webdriver.Remote("http://127.0.0.1:4723", options=options)

    try:

        log("CLEANUP", "Opening Recent Apps")
        driver.press_keycode(187)
        time.sleep(3)

        close_all_locators = [

            # Samsung tablets
            "//*[@text='CLOSE ALL']",
            "//*[contains(@text,'CLOSE ALL')]",

            # Samsung mixed case
            "//*[@text='Close all']",
            "//*[contains(@text,'Close all')]",

            # Other Android devices
            "//*[@text='CLEAR ALL']",
            "//*[contains(@text,'CLEAR ALL')]",

            "//*[@text='Clear all']",
            "//*[contains(@text,'Clear all')]",

            # Accessibility
            "//*[contains(@content-desc,'CLOSE ALL')]",
            "//*[contains(@content-desc,'Close all')]",
            "//*[contains(@content-desc,'CLEAR ALL')]",
            "//*[contains(@content-desc,'Clear all')]",
        ]

        button_clicked = False

        for locator in close_all_locators:

            try:

                clear_all = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.XPATH, locator))
                )

                clear_all.click()

                log("CLEANUP", "Recent apps cleared")

                button_clicked = True

                break

            except:
                continue

        if not button_clicked:
            log("CLEANUP", "No recent apps to clear")

        # print("Recent apps cleared via UI.")
        # log("CLEANUP", "Recent apps cleared")

    except Exception as e:
        print(f"No recent apps to clear or UI not found: {e}")

    finally:
        driver.quit()


# ========== Save to Excel ==========
timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
# excel_path = f"C:/Users/SVC-Systems-TestPC/OneDrive - Medtronic PLC/ProtocolApp_Versions_Pulling_{timestamp}.xlsx"
excel_path = f"C:/Users/mandat3/OneDrive - Medtronic PLC/Desktop/Appium/App_Versions_Pulling/App_Versions_Pulling_{timestamp}.xlsx"
df = pd.DataFrame(app_data)
df.to_excel(excel_path, index=False)
print(f"Versions saved to: {excel_path}")

# ========== Update Word Document ==========

# word_path = r"C:\Users\mandat3\OneDrive - Medtronic PLC\Desktop\Appium\App Ver Auto Fill\NDHF1500-221664_Jan_2026_Protocol.docx"
word_path = r"C:\Users\mandat3\OneDrive - Medtronic PLC\Desktop\Appium\App Ver Auto Fill\NDHF1500-221664_Jan_2026_Protocol.docx"
# word_path = r"C:\Users\SVC-Systems-TestPC\OneDrive - Medtronic PLC\Protocol\NDHF1500-221664_CP_Latest_Protocol.docx"
# output_word_path = word_path.replace(".docx", "_Updated.docx")
output_word_path = word_path

if not os.path.exists(word_path):
    print(f"Word document not found at: {word_path}")
else:
    if os.path.exists(output_word_path):
        # print("Loading existing updated document...")
        log("DOCUMENT", "Loading protocol document")
        doc = Document(output_word_path)
    else:
        # print("Loading original Word document...")
        log("DOCUMENT", "Loading protocol document")
        doc = Document(word_path)

    s2_versions = {row["App"].strip().lower(): row["Version"] for _, row in df.iterrows()}

    for table in doc.tables:
        if len(table.rows) < 3 or len(table.rows[1].cells) < 7:
            continue

        second_header_row = table.rows[1]
        # header_labels = [cell.text.strip() for cell in second_header_row.cells]
        header_labels = [cell.text.strip().lower() for cell in second_header_row.cells]
        # print("Detected Headers:", header_labels)

        # Find all S2 columns
        s2_indices = [i for i, label in enumerate(header_labels) if label == "S2" or "s2" in label]

        print("S2 Indexes:", s2_indices)

        if len(s2_indices) < 3:
            print("Expected 3 S2 columns (Dec, Jan, Remarks).")
            continue

        # Column mapping
        Dec_s2_col_Prev = s2_indices[0]
        Jan_s2_col_Cur = s2_indices[1]
        remarks_s2_col = s2_indices[2]

        # print(f"Updating Jan S2 column at index {Jan_s2_col_Cur} and Remarks S2 at {remarks_s2_col}")

        for row in table.rows[2:]:
            if len(row.cells) <= remarks_s2_col:
                continue

            app_cell_text = row.cells[1].text.strip().lower()
            matched_key = next((key for key in s2_versions if key in app_cell_text), None)
            version = s2_versions.get(matched_key, "N/A") if matched_key else "N/A"

            # Update Jan S2 version
            Jan_cell = row.cells[Jan_s2_col_Cur]
            Jan_cell.text = version
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
            Jan_cell._tc.get_or_add_tcPr().append(shading_elm)
            # print(f"{app_cell_text} → {version}")
            log(
                "DOCUMENT",
                f"{app_cell_text:<30} → {version:<15}"
            )

            # Compare with Dec value for remarks
            Dec_val = row.cells[Dec_s2_col_Prev].text.strip()
            Jan_val = version.strip()

            if not Dec_val or not Jan_val or "N/A" in (Dec_val, Jan_val):
                remark = "No"
            else:
                remark = "Yes" if Dec_val != Jan_val else "No"

            row.cells[remarks_s2_col].text = remark
            # print(f"Remark for s2 → {remark}")
            log(
                "DOCUMENT",
                f"{app_cell_text:<30} → {version:<15}"
            )

        break  # Only update first matched table

    doc.save(output_word_path)

    log("DOCUMENT", "Word document updated successfully")
    log("DOCUMENT", output_word_path)

    # --------------------------------------------------
    # CLEANUP BEFORE SOFTWARE UPDATE
    # --------------------------------------------------

    clear_recent_apps_ui(device_serial)

    # --------------------------------------------------
    # SOFTWARE UPDATE CHECK
    # --------------------------------------------------

    check_software_update()

    # --------------------------------------------------
    # FINAL CLEANUP
    # --------------------------------------------------

    clear_recent_apps_ui(device_serial)

    log("RESULT", "S2 extraction completed successfully")



