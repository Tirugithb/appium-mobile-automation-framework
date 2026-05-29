import os
import time
import subprocess
from datetime import datetime

from appium import webdriver
from appium.options.android import UiAutomator2Options
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC

from docx import Document
from docx.shared import RGBColor
from packaging import version
import sys

sys.stdout.reconfigure(encoding="utf-8")

# add project root
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))

# --------------------------------------------------
# CORE IMPORTS
# --------------------------------------------------
from core.device_registry import is_device_allowed, get_device_name
from core.adb_utils import (
    get_connected_device_serial,
    is_device_locked,
    unlock_device_with_password
)
from core.config_loader import load_device_password

# --------------------------------------------------
# DISPLAY SERIAL MAPPING
# --------------------------------------------------

DISPLAY_SERIAL_MAP = {
    "22a9e8a977239f73": "R52M10M43NV"
}

# --------------------------------------------------
# LOGGING
# --------------------------------------------------
def log(section, message):
    ts = datetime.now().strftime("%H:%M:%S")
    print(f"[{ts}] [{section:<14}] {message}")


# --------------------------------------------------
# DEVICE RESOLUTION & UNLOCK
# --------------------------------------------------
device_serial = os.environ.get("DEVICE_SERIAL") or get_connected_device_serial()
display_serial = DISPLAY_SERIAL_MAP.get(
    device_serial,
    device_serial
)

if not is_device_allowed(device_serial):
    raise RuntimeError(f"Unauthorized S2 device connected: {device_serial}")

device_name = get_device_name(device_serial)
# log("DEVICE", f"Authorized device detected: {device_name} ({device_serial})")
log(
    "DEVICE",
    f"Authorized device detected: {device_name} ({display_serial})"
)
if is_device_locked():
    unlock_device_with_password(load_device_password())
else:
    log("DEVICE", "Device already unlocked")
    
    
# --------------------------------------------------
# START APPIUM
# --------------------------------------------------
from appiumManager.appium_manager import AppiumManager
AppiumManager.start()

# --------------------------------------------------
# CONSTANTS
# --------------------------------------------------
DOC_PATH = r"C:\Users\mandat3\OneDrive - Medtronic PLC\Desktop\Appium\App Ver Auto Fill\NDHF1500-221664_Jan_2026_Protocol.docx"
# DOC_PATH = r"C:\Users\SVC-Systems-TestPC\OneDrive - Medtronic PLC\Protocol\NDHF1500-221664_CP_Latest_Protocol.docx"

APP_NAME_MAPPING = {
    "DBS Clinician Application (A610)": "DBS",
    "DBS Patient Demo (A620)": "DBS Patient Demo",
    "SureTune 4 Connector (A90400)": "SureTune™ 4 Connector",
    "SynchroMed Clinician Application (A810)": "SynchroMed™",
    "Restore Clinician Application (A71100)": "Restore",
    "Vanta Clinician Application (A71200)": "Vanta",
    "Stim Trialing Clinician Application (A71300)": "Stim Trialing",
    "Intellis Clinician Application (A710)": "Intellis",
    "Inceptiv Clinician Application (A71400)": "Inceptiv",
    "Altaviva Clinician Application (P7850N)": "Altaviva Clinician",
    "Recharger Application (A90300)": "Recharger Application",
    "Communication Manager (A901)": "Medtronic Communication Manager",
    "PDS (A902)": "Patient Data Service"
}


# --------------------------------------------------
# DEVICE DETECTION
# --------------------------------------------------
def get_all_connected_devices():
    output = subprocess.check_output("adb devices", shell=True).decode().strip()
    return [line.split()[0] for line in output.splitlines()[1:] if "device" in line]


def get_device_model(serial):
    try:
        return subprocess.check_output(
            f"adb -s {serial} shell getprop ro.product.model", shell=True
        ).decode().strip()
    except:
        return None


def wait_for_device_model(target_model, prompt="Connect device"):
    log("DEVICE", f"{prompt} ({target_model})")
    while True:
        for serial in get_all_connected_devices():
            if get_device_model(serial) == target_model:
                log("DEVICE", f"{target_model} connected: {serial}")
                return serial
        time.sleep(2)


# --------------------------------------------------
# HUB HELPERS
# --------------------------------------------------
def force_stop_and_launch_hub(serial):
    log("HUB", f"Launching Hub on {serial}")

    subprocess.run(
        f"adb -s {serial} shell am force-stop com.airwatch.androidagent",
        shell=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL
    )

    subprocess.run(
        f"adb -s {serial} shell monkey -p com.airwatch.androidagent "
        f"-c android.intent.category.LAUNCHER 1",
        shell=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL
    )

    time.sleep(5)


def scroll_down(driver):
    driver.execute_script("mobile: scrollGesture", {
        "left": 100,
        "top": 300,
        "width": 600,
        "height": 800,
        "direction": "down",
        "percent": 0.85
    })
    time.sleep(1)

def click_managed_apps(driver):

    log("HUB", "Searching for Managed Apps")

    for i in range(8):   # try scrolling 8 times

        elements = driver.find_elements(
            By.XPATH,
            "//android.widget.TextView[@text='Managed Apps']"
        )

        if elements:
            elements[0].click()
            log("HUB", "Managed Apps opened")
            return

        log("HUB", f"Managed Apps not visible — scrolling ({i+1}/8)")
        scroll_down(driver)

    raise Exception("Managed Apps not found after scrolling")

# --------------------------------------------------
# WORD DOC UPDATE
# --------------------------------------------------
def fill_group_id_below_aw_group(doc, group_id):
    found = False
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip().startswith("S2:"):
                        para.clear()
                        run1 = para.add_run("S2: ")
                        run2 = para.add_run(group_id)
                        run2.font.color.rgb = RGBColor(0, 128, 0)
                        found = True
    if not found:
        print("Could not find 'S2:' line to fill Group ID.")


def fill_into_word_table(app_versions, group_id):
    target_col_index = 2   # S2 column
    doc = Document(DOC_PATH)
    for table in doc.tables:
        for row in table.rows[1:]:
            if len(row.cells) < 6:
                continue
            app_name_in_doc = row.cells[1].text.strip()
            tablet_app_name = APP_NAME_MAPPING.get(app_name_in_doc)
            if not tablet_app_name:
                continue
            version = app_versions.get(tablet_app_name, "N/A")
            cell = row.cells[target_col_index]
            cell.text = version
            if cell.paragraphs and cell.paragraphs[0].runs:
                run = cell.paragraphs[0].runs[0]
                if version != "N/A":
                    run.font.color.rgb = RGBColor(0, 128, 0)
                else:
                    run.font.color.rgb = RGBColor(255, 0, 0)
    fill_group_id_below_aw_group(doc, group_id)
    doc.save(DOC_PATH)
    print(f"S2 versions filled and saved to {DOC_PATH}")


# --------------------------------------------------
# EXTRACTION LOGIC
# --------------------------------------------------
def extract_uat_group_and_apps(serial):
    force_stop_and_launch_hub(serial)
    time.sleep(5)

    options = UiAutomator2Options()
    options.platform_name = "Android"
    options.device_name = serial
    options.automation_name = "UiAutomator2"
    options.app_package = "com.airwatch.androidagent"
    options.app_activity = "com.airwatch.agent.Hub.hostactivity.HostActivity"

    options.no_reset = True
    options.dont_stop_app_on_reset = True
    options.skip_server_installation = True
    options.allow_running_instrumentation = True
    options.ignore_hidden_api_policy_error = True

    options.uiautomator2_server_launch_timeout = 180000
    options.uiautomator2_server_install_timeout = 180000
    options.adb_exec_timeout = 180000
    options.android_install_timeout = 180000

    driver = webdriver.Remote("http://127.0.0.1:4723", options=options)

    wait = WebDriverWait(driver, 30)

    # REQUIRED UI READY CHECK
    wait.until(
        EC.presence_of_element_located(
            (By.ID, "com.airwatch.androidagent:id/user_initials_tv")
        )
    )
    log("HUB", "Hub UI ready")

    app_versions = {}
    uat = group_id = "N/A"

    try:
        wait.until(EC.element_to_be_clickable((By.ID, "com.airwatch.androidagent:id/user_initials_tv"))).click()
        wait.until(EC.element_to_be_clickable((By.ID, "com.airwatch.androidagent:id/device_tv"))).click()
        wait.until(EC.element_to_be_clickable((By.XPATH, "//android.widget.TextView[@text='Enrollment']"))).click()

        values = driver.find_elements(By.ID, "com.airwatch.androidagent:id/listview_value")
        if len(values) >= 2:
            uat = values[0].text.strip()
            group_id = values[1].text.strip()
            log("HUB", f"UAT: {uat} | Group ID: {group_id}")

        driver.find_element(By.XPATH, "//android.widget.ImageButton[@content-desc='Back']").click()
        time.sleep(2)

        # wait.until(EC.element_to_be_clickable((By.XPATH, "//android.widget.TextView[@text='Managed Apps']"))).click()
        click_managed_apps(driver)

        # --------------------------------------------------
        # WAIT FOR MANAGED APPS TO LOAD
        # --------------------------------------------------

        log("HUB", "Waiting for Managed Apps list to load")

        time.sleep(8)

        last_seen = ""
        stable_count = 0

        while True:

            app_names = driver.find_elements(
                By.ID,
                "com.airwatch.androidagent:id/app_name"
            )

            app_versions_ui = driver.find_elements(
                By.ID,
                "com.airwatch.androidagent:id/app_version"
            )

            new_found = False

            for name_el, ver_el in zip(app_names, app_versions_ui):

                try:

                    name = name_el.text.strip()
                    ver = ver_el.text.strip()

                    if name and name not in app_versions:
                        app_versions[name] = ver

                        log("HUB APPS", f"{name} - {ver}")

                        new_found = True

                except:
                    pass

            last = list(app_versions.keys())[-1] if app_versions else ""

            if last == last_seen:
                stable_count += 1
            else:
                stable_count = 0
                last_seen = last

            if stable_count >= 2:
                break

            scroll_down(driver)

    finally:
        driver.quit()
        log("HUB", "Appium session closed")

    return uat, group_id, app_versions


# --------------------------------------------------
# CLEANUP
# --------------------------------------------------
def clear_recent_apps_ui(serial):
    options = UiAutomator2Options()
    options.platform_name = "Android"
    options.device_name = serial
    options.no_reset = True  # IMPORTANT
    # DO NOT set app_package / app_activity

    # driver = webdriver.Remote("http://localhost:4723", options=options)
    driver = webdriver.Remote("http://127.0.0.1:4723", options=options)

    try:

        log("CLEANUP", "Opening Recent Apps")
        driver.press_keycode(187)
        time.sleep(3)

        close_all_locators = [

            # Samsung tablets
            "//*[@text='CLOSE ALL']",
            "//*[contains(@text,'CLOSE ALL')]",

            # Samsung mixed case
            "//*[@text='Close all']",
            "//*[contains(@text,'Close all')]",

            # Other Android devices
            "//*[@text='CLEAR ALL']",
            "//*[contains(@text,'CLEAR ALL')]",

            "//*[@text='Clear all']",
            "//*[contains(@text,'Clear all')]",

            # Accessibility
            "//*[contains(@content-desc,'CLOSE ALL')]",
            "//*[contains(@content-desc,'Close all')]",
            "//*[contains(@content-desc,'CLEAR ALL')]",
            "//*[contains(@content-desc,'Clear all')]",
        ]

        button_clicked = False

        for locator in close_all_locators:

            try:

                clear_all = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.XPATH, locator))
                )

                clear_all.click()

                log("CLEANUP", "Recent apps cleared")

                button_clicked = True

                break

            except:
                continue

        if not button_clicked:
            log("CLEANUP", "No recent apps to clear")

        # print("Recent apps cleared via UI.")
        # log("CLEANUP", "Recent apps cleared")

    except Exception as e:
        print(f"No recent apps to clear or UI not found: {e}")

    finally:
        driver.quit()


# --------------------------------------------------
# MAIN
# --------------------------------------------------
if __name__ == "__main__":
    log("INIT", "Starting extraction for S2")

    device = device_serial

    uat, group, versions = extract_uat_group_and_apps(device)

    fill_into_word_table(versions, group)

    clear_recent_apps_ui(device)

    log("RESULT", "Extraction complete for S2")
