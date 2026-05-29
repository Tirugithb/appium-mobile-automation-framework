from appium import webdriver
from appium.options.android import UiAutomator2Options
# from selenium.webdriver.common.by import By
from appium.webdriver.common.appiumby import AppiumBy as By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver import ActionChains
from selenium.webdriver.common.action_chains import ActionBuilder
from selenium.webdriver.common.actions import interaction
from selenium.webdriver.common.actions.pointer_input import PointerInput
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from datetime import datetime
import time
import pandas as pd
import os
import subprocess
import sys

sys.stdout.reconfigure(encoding="utf-8")

# add project root
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..")))

# --------------------------------------------------
# CORE IMPORTS
# --------------------------------------------------
from core.device_registry import is_device_allowed, get_device_name
from core.adb_utils import (
    get_connected_device_serial,
    is_device_locked,
    unlock_device_with_password
)
from core.config_loader import load_device_password


# --------------------------------------------------
# LOGGING
# --------------------------------------------------
def log(section, message):
    ts = datetime.now().strftime("%H:%M:%S")
    print(f"[{ts}] [{section:<14}] {message}")


# --------------------------------------------------
# DEVICE RESOLUTION & UNLOCK
# --------------------------------------------------
device_serial = os.environ.get("DEVICE_SERIAL") or get_connected_device_serial()

if not is_device_allowed(device_serial):
    raise RuntimeError(f"Unauthorized A4P device connected: {device_serial}")

device_name = get_device_name(device_serial)
log("DEVICE", f"Authorized device detected: {device_name} ({device_serial})")

if is_device_locked():
    unlock_device_with_password(load_device_password())
else:
    log("DEVICE", "Device already unlocked")

# --------------------------------------------------
# START APPIUM
# --------------------------------------------------
from appiumManager.appium_manager import AppiumManager

AppiumManager.start()


def force_stop_and_launch_settings(serial):
    # print(f"Launching Settings on {serial}")
    log("SYSTEM UPDATE", f"Launching Settings on {serial}")

    subprocess.run(
        f"adb -s {serial} shell am force-stop com.android.settings",
        shell=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL
    )

    time.sleep(2)

    subprocess.run(
        f"adb -s {serial} shell monkey -p com.android.settings "
        f"-c android.intent.category.LAUNCHER 1",
        shell=True,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL
    )

    time.sleep(5)


# ========== Appium Setup ==========
def get_settings_driver(serial):
    options = UiAutomator2Options()
    options.platform_name = "Android"
    options.device_name = serial
    options.udid = serial
    options.automation_name = "UiAutomator2"
    options.app_package = "com.android.settings"
    options.app_activity = "com.android.settings.Settings"
    options.no_reset = True
    # ADD – required for Samsung / slow instrumentation
    options.uiautomator2_server_launch_timeout = 180000
    options.uiautomator2_server_install_timeout = 180000
    options.adb_exec_timeout = 180000
    options.android_install_timeout = 180000
    # ADD – because Settings is launched via adb monkey
    options.allow_running_instrumentation = True
    options.ignore_hidden_api_policy_error = True

    # return webdriver.Remote("http://localhost:4723", options=options)
    return webdriver.Remote("http://127.0.0.1:4723", options=options)


# ========== Helper Functions ==========

def swipe_up_left_side(driver):
    size = driver.get_window_size()
    start_x = size['width'] * 0.2
    start_y = size['height'] * 0.8
    end_y = size['height'] * 0.2
    actions = ActionChains(driver)
    pointer = PointerInput(interaction.POINTER_TOUCH, "finger")
    actions.w3c_actions = ActionBuilder(driver, mouse=pointer)
    actions.w3c_actions.pointer_action.move_to_location(start_x, start_y)
    actions.w3c_actions.pointer_action.pointer_down()
    actions.w3c_actions.pointer_action.move_to_location(start_x, end_y)
    actions.w3c_actions.pointer_action.release()
    actions.perform()
    time.sleep(1)


def swipe_up_in_apps_list(driver):
    size = driver.get_window_size()
    start_x = size['width'] * 0.8
    start_y = size['height'] * 0.8
    end_y = size['height'] * 0.3
    actions = ActionChains(driver)
    pointer = PointerInput(interaction.POINTER_TOUCH, "finger")
    actions.w3c_actions = ActionBuilder(driver, mouse=pointer)
    actions.w3c_actions.pointer_action.move_to_location(start_x, start_y)
    actions.w3c_actions.pointer_action.pointer_down()
    actions.w3c_actions.pointer_action.pause(0.5)
    actions.w3c_actions.pointer_action.move_to_location(start_x, end_y)
    actions.w3c_actions.pointer_action.pause(0.2)
    actions.w3c_actions.pointer_action.release()
    actions.perform()
    time.sleep(0.3)


def wait_for_apps_list_loaded(driver, timeout=10):
    try:
        WebDriverWait(driver, timeout).until(
            EC.presence_of_element_located((By.XPATH, '//android.widget.TextView[contains(@text, "Apps")]'))
        )
    except:
        print("Apps list did not load.")


def find_and_click_apps(driver):
    try:
        apps_option = WebDriverWait(driver, 20).until(
            EC.presence_of_element_located((By.XPATH, '//android.widget.TextView[@text="Apps"]'))
        )
        apps_option.click()
    except:
        scroll_to_top_left(driver)
        find_and_click_apps(driver)


def scroll_until_app_found(driver, app_name, max_swipes=4):
    swipes = 0
    previous_page_source = ""
    while swipes < max_swipes:
        try:
            wait_for_apps_list_loaded(driver)
            app_option = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.XPATH, f'//android.widget.TextView[contains(@text, "{app_name}")]'))
            )
            app_option.click()
            return True
        except:
            swipe_up_in_apps_list(driver)
            time.sleep(2)
            current_page_source = driver.page_source
            if current_page_source == previous_page_source:
                break
            previous_page_source = current_page_source
            swipes += 1
    return False


def scroll_up_to_version_info(driver):
    size = driver.get_window_size()
    start_x = size['width'] / 2
    start_y = size['height'] * 0.7
    end_y = size['height'] * 0.3
    actions = ActionChains(driver)
    pointer = PointerInput(interaction.POINTER_TOUCH, "finger")
    actions.w3c_actions = ActionBuilder(driver, mouse=pointer)
    actions.w3c_actions.pointer_action.move_to_location(start_x, start_y)
    actions.w3c_actions.pointer_action.pointer_down()
    actions.w3c_actions.pointer_action.pause(0.2)
    actions.w3c_actions.pointer_action.move_to_location(start_x, end_y)
    actions.w3c_actions.pointer_action.release()
    actions.perform()
    time.sleep(1.5)
    try:
        version_element = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//android.widget.TextView[contains(@text, "Version")]'))
        )
        return version_element.text
    except:
        return None


def navigate_back(driver):
    try:
        back_button = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//android.widget.ImageButton[@content-desc="Navigate up"]'))
        )
        back_button.click()
        time.sleep(2)
    except:
        print("Back button not found.")


def scroll_to_top_left(driver, max_swipes=3):
    for _ in range(max_swipes):
        swipe_up_left_side(driver)
        time.sleep(2)


# ========== Start Process ==========
force_stop_and_launch_settings(device_serial)
driver = get_settings_driver(device_serial)

scroll_to_top_left(driver)
find_and_click_apps(driver)

apps_to_find = ["Adobe Acrobat", "Brother Print Service Plugin", "Calculator", "Camera", "Canon Print Service",
                "Clock", "Connector", "Content", "Epson iPrint", "Gallery", "Google Play Store", "Hub", "iProjection",
                "Knox E-FOTA", "Mopria Print Service", "My Files", "Ricoh Smart Device Connector", "Teams", "Web",
                "Xerox Print Service"]
app_data = []

for app_name in apps_to_find:
    if scroll_until_app_found(driver, app_name):
        version_info = scroll_up_to_version_info(driver)
        version_number = version_info.replace("Version ", "") if version_info else "N/A"  # Changed to N/A to No
        # print(f"{app_name} : {version_number}")
        log("SETTINGS APPS", f"{app_name:<30} : {version_number}")
        app_data.append({"App": app_name, "Version": version_number})
        navigate_back(driver)
    else:
        # print(f"{app_name} not found.")
        log("SETTINGS APPS", f"{app_name:<30} : N/A")  # Changed to N/A to No
        app_data.append({"App": app_name, "Version": "N/A"})  # Changed to N/A to No

driver.quit()


# --------------------------------------------------
# CLEANUP
# --------------------------------------------------
def clear_recent_apps_ui(serial):
    options = UiAutomator2Options()
    options.platform_name = "Android"
    options.device_name = serial
    options.no_reset = True  # IMPORTANT
    # DO NOT set app_package / app_activity

    # driver = webdriver.Remote("http://localhost:4723", options=options)
    driver = webdriver.Remote("http://127.0.0.1:4723", options=options)

    try:
        # print("Opening Recent Apps…")
        log("CLEANUP", "Opening Recent Apps")
        driver.press_keycode(187)
        time.sleep(3)

        clear_all = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable(
                (By.XPATH, "//*[@text='Close all' or @text='Clear all' or contains(@content-desc,'Clear')]")))
        clear_all.click()
        # print("Recent apps cleared via UI.")
        log("CLEANUP", "Recent apps cleared")

    except Exception as e:
        print(f"No recent apps to clear or UI not found: {e}")

    finally:
        driver.quit()


# ========== Save to Excel ==========
timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
excel_path = f"C:/Users/SVC-Systems-TestPC/OneDrive - Medtronic PLC/ProtocolApp_Versions_Pulling_{timestamp}.xlsx"
# excel_path = f"C:/Users/mandat3/OneDrive - Medtronic PLC/Desktop/Appium/App_Versions_Pulling/App_Versions_Pulling_{timestamp}.xlsx"
# excel_path = f"C:/Users/mohami32/OneDrive - Medtronic PLC/Desktop/Appium/App_Versions_Pulling/App_Versions_Pulling_{timestamp}.xlsx"

df = pd.DataFrame(app_data)
df.to_excel(excel_path, index=False)
print(f"Versions saved to: {excel_path}")

# ========== Update Word Document ==========

# word_path = r"C:\Users\mandat3\OneDrive - Medtronic PLC\Desktop\Appium\App Ver Auto Fill\NDHF1500-221664_Jan_2026_Protocol.docx"
# word_path = r"C:\Users\mohami32\OneDrive - Medtronic PLC\Desktop\Appium\App Ver Auto Fill\NDHF1500-221664_Jan_2026_Protocol.docx"
word_path = r"C:\Users\SVC-Systems-TestPC\OneDrive - Medtronic PLC\Protocol\NDHF1500-221664_CP_Latest_Protocol.docx"
# output_word_path = word_path.replace(".docx", "_Updated.docx")
output_word_path = word_path

if not os.path.exists(word_path):
    print(f"Word document not found at: {word_path}")
else:
    if os.path.exists(output_word_path):
        # print("Loading existing updated document...")
        log("DOCUMENT", "Loading protocol document")
        doc = Document(output_word_path)
    else:
        # print("Loading original Word document...")
        log("DOCUMENT", "Loading protocol document")
        doc = Document(word_path)

    a4p_versions = {row["App"].strip().lower(): row["Version"] for _, row in df.iterrows()}

    for table in doc.tables:
        if len(table.rows) < 3 or len(table.rows[1].cells) < 7:
            continue

        second_header_row = table.rows[1]
        header_labels = [cell.text.strip() for cell in second_header_row.cells]

        # Find Dec and Jan A4P columns
        a4p_indices = [i for i, label in enumerate(header_labels) if label == "A4P"]
        if len(a4p_indices) < 2:
            print("Could not find both Dec and Jan A4P columns.")
            continue

        Dec_a4p_col_Prev = a4p_indices[0]
        Jan_a4p_col_Cur = a4p_indices[1]

        # FIX: Find Remarks A4P column safely by name
        a4p_indices = [i for i, label in enumerate(header_labels) if label == "A4P"]
        if len(a4p_indices) < 3:
            print("Expected at least 3 A4P columns (Dec, Jan, Remarks).")
            continue

        Dec_a4p_col_Prev = a4p_indices[0]
        Jan_a4p_col_Cur = a4p_indices[1]
        remarks_a4p_col = a4p_indices[2]

        # print(f"Updating Jan A4P column at index {Jan_a4p_col_Cur} and Remarks A4P at {remarks_a4p_col}")

        # print(f"Updating Jan A4P column at index {Jan_a4p_col_Cur} and Remarks A4P at {remarks_a4p_col}")

        for row in table.rows[2:]:
            if len(row.cells) <= remarks_a4p_col:
                continue

            app_cell_text = row.cells[1].text.strip().lower()
            matched_key = next((key for key in a4p_versions if key in app_cell_text), None)
            version = a4p_versions.get(matched_key, "N/A") if matched_key else "N/A"  # Changed N/A to No

            # Update Jan A4P version
            Jan_cell = row.cells[Jan_a4p_col_Cur]
            Jan_cell.text = version
            shading_elm = parse_xml(r'<w:shd {} w:fill="D9EAD3"/>'.format(nsdecls('w')))
            Jan_cell._tc.get_or_add_tcPr().append(shading_elm)
            # print(f"{app_cell_text} → {version}")
            log(
                "DOCUMENT",
                f"{app_cell_text:<30} → {version:<15}"
            )

            # Compare with Dec value for remarks
            Dec_val = row.cells[Dec_a4p_col_Prev].text.strip()
            Jan_val = version.strip()

            if not Dec_val or not Jan_val or "N/A" in (Dec_val, Jan_val):  # Changed N/A to No
                remark = "No"
            else:
                remark = "Yes" if Dec_val != Jan_val else "No"

            row.cells[remarks_a4p_col].text = remark
            # print(f"Remark for A4P → {remark}")
            log(
                "DOCUMENT",
                f"{app_cell_text:<30} → {version:<15}"
            )

        break  # Only update first matched table

    doc.save(output_word_path)
    # print(f"Word document updated: {output_word_path}")
    log("DOCUMENT", "Word document updated successfully")
    log("DOCUMENT", output_word_path)

    clear_recent_apps_ui(device_serial)
    log("RESULT", "A4P extraction completed successfully")



